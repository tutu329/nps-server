from docx import Document
from docx.oxml.xmlchemy import BaseOxmlElement, ZeroOrOne, ZeroOrMore, OxmlElement
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.oxml.ns import qn
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.enum.section import WD_SECTION_START
from docx.enum.section import WD_ORIENTATION

from PIL import Image
import matplotlib.pyplot as plt
import os

import Global
Global.init()

class Word():
    def __init__(self, in_filename, in_id="", in_domain_name=""):
        self.docx_filename = Global.get("path") + "static/docx/" + in_filename + "-openid(" + in_id + ").docx"
        self.pic_filepath = Global.get("path") + "static/pic/"
        self._docx_full_url = in_domain_name + "/static/docx/" + in_filename + "-openid(" + in_id + ").docx"
        self._doc = Document()
        self._id = in_id
        self._temp_files = []

        # 设置Title样式
        self.create_style(style_name="报告标题", font_name="黑体", font_bold=True, style_type=1, font_size=18, font_color=[0x00, 0x00, 0x00])

        # 设置正文样式
        self._doc.styles['Normal'].font.size = Pt(12)
        self._doc.styles['Normal'].font.bold = False
        self._doc.styles['Normal'].font.italic = False
        self._doc.styles['Normal'].font.name = 'Times New Roman'
        self._doc.styles['Normal'].font.color.rgb = RGBColor(0, 0, 0)
        self._doc.styles['Normal'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # 页脚
        # self._doc.sections[0].footer.paragraphs[0].text = "这是第1节页脚"
        # self._doc.sections[0].footer.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        # 纸张大小
        # doc.sections[0].page_width = Cm(21)
        # doc.sections[0].page_height = Cm(29.7)
        # 页边距
        # doc.sections[0].top_margin = Cm(2)  # 修改上页边距
        # doc.sections[0].bottom_margin = Cm(2)  # 修改上页边距
        # doc.sections[0].left_margin = Cm(2)  # 修改上页边距
        # doc.sections[0].right_margin = Cm(2)  # 修改上页边距
        # 装订线
        # doc.sections[0].gutter = Cm(2)  # 设置装订线为2cm
        # 页眉
        # doc.sections[0].header.paragraphs[0].text = "这是第1节页眉"
        # doc.sections[0].header_distance = Cm(10)
        # 页脚

        # 奇偶页页眉页脚
        # doc.settings.odd_and_even_pages_header_footer = True

        # 表格内容居中
        # table.cell(r, c).paragraphs[0].paragraph_format.alignment = WD_TABLE_ALIGNMENT.CENTER

        # 字号对应
        # 初号 = 42        磅
        # 小初 = 36        磅
        # 一号 = 26        磅
        # 小一 = 24        磅
        # 二号 = 22        磅
        # 小二 = 18        磅
        # 三号 = 16        磅
        # 小三 = 15        磅
        # 四号 = 14        磅
        # 小四 = 12        磅
        # 五号 = 10.5      磅
        # 小五 = 9         磅
        # 六号 = 7.5       磅
        # 小六 = 6.5       磅
        # 七号 = 5.5       磅
        # 八号 = 5         磅

        # 设置表格样式
        # self._doc.styles['"Light Grid Accent 3"'].font.size = Pt(10.5)
        # self._doc.styles['"Light Grid Accent 3"'].font.name = 'Times New Roman'
        # self._doc.styles['"Light Grid Accent 3"'].font.color.rgb = RGBColor(0, 0, 0)
        # self._doc.styles['"Light Grid Accent 3"'].element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def heading(self, in_str, in_level):
        # 添加标题
        para = self._doc.add_heading(level=in_level)
        run = para.add_run(in_str)


        # 设置style
        if in_level==1 :
            run.font.size=Pt(14)
            run.font.bold=True
            run.font.name='黑体'
            run.font.color.rgb=RGBColor(0,0,0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            # 段前段后
            para.paragraph_format.space_before = Pt(12)
            para.paragraph_format.space_after = Pt(6)
        elif in_level==2 :
            run.font.size=Pt(14)
            run.font.bold=True
            run.font.name='宋体'
            run.font.color.rgb=RGBColor(0,0,0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            para.paragraph_format.space_before = Pt(6)
            para.paragraph_format.space_after = Pt(6)
        elif in_level==3 :
            run.font.size=Pt(14)
            run.font.name='宋体'
            run.font.color.rgb=RGBColor(0,0,0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = Pt(3)
        elif in_level==4 :
            run.font.size=Pt(14)
            run.font.name='宋体'
            run.font.color.rgb=RGBColor(0,0,0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0
        elif in_level==5 :
            run.font.size=Pt(14)
            run.font.name='宋体'
            run.font.color.rgb=RGBColor(0,0,0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0
        else:
            run.font.size = Pt(14)
            run.font.name = '宋体'
            run.font.color.rgb = RGBColor(0, 0, 0)
            run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
            para.paragraph_format.space_before = 0
            para.paragraph_format.space_after = 0
        return self
        # run.font.size = Pt(14)
        # run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # t_heading.style.font.size = Pt(20)
        # t_heading.style.font.name = '宋体'
        # t_heading.style._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

        # head = document.add_heading(0)
        # run = head.add_run('需要制作的文档')
        # run.font.size = Pt(24)

        # obj.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        # 需要导入：from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

        # t_heading._p.pPr.pStyle.set(qn('w:val'), u'4FDD')

        # t_title.font.name = '宋体'
        # # 设置中文字体
        # t_title._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # # 设置字体颜色
        # t_title.font.color.rgb = RGBColor(255, 55, 55)  # 红色
        # # 设置字体大小
        # t_title.font.size = Pt(30)
        # # 设置下划线
        # t_title.font.underline = True
        # # 设置删除线
        # t_title.font.strike = True
        # # 设置加粗
        # t_title.bold = True
        # # 设置斜体
        # t_title.italic = True

    def para(self, in_str, in_style="Normal", in_center=False, in_indent=True, in_line_space=1.5, in_space_before=0, in_space_after=0):
        t_list = in_str.split("\n")

        for t_str in t_list:
            t_para = self._doc.add_paragraph(t_str, style=in_style)
            t_para.paragraph_format.line_spacing = in_line_space  # 设置该段落，行间距为1.5倍，也可以像上面设默认值那样用Pt单位来设置
            if in_indent:
                t_para.paragraph_format.first_line_indent = Pt(14) * 2  # 段落缩进0.5英寸，我还是习惯设置2字符 值为：406400

            if in_center :
                t_para.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            # 段前段后
            t_para.paragraph_format.space_before = in_space_before
            t_para.paragraph_format.space_after = in_space_after

        # if in_size!=0 :
        #     t_para.style.font.size = in_size

        # t_para.paragraph_format.left_line_indent = Inches(0.5)  # 设置左缩进0.5英寸。一般用不到
        # p1.paragraph_format.right_line_indent = Inches(0.5)  # 设置右缩进0.5英寸，一般用不到
        # p1.paragraph_format.keep_together = False  # 段前分页
        # p1.paragraph_format.keep_with_next = False  # 与下段同页
        # p1.paragraph_format.page_break_before = True  # 段中不分页
        # p1.paragraph_format.widow_control = False  # 孤行控制

        return t_para

    def table(self, in_table_head_list, in_table_data_list, in_title="", in_style="Light Grid Accent 3"):
        t_cols = len(in_table_head_list)
        t_rows = len(in_table_data_list)

        if in_title!="":
            self.para(in_title, in_style="Normal", in_center=True, in_indent=False, in_line_space=1.0)

        # 添加表头
        t_table = self._doc.add_table(rows=1, cols=t_cols, style=in_style)

        for i in range(t_cols) :
            # print("i=",i)
            # t_table.rows[0].cells[i].text = in_table_head_list[i]
            # t_table.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # t_table.rows[0].cells[i].vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            t_table.rows[0].cells[i].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            run = t_table.rows[0].cells[i].paragraphs[0].add_run(in_table_head_list[i])
            run.font.name = '宋体'
            run.font.size = Pt(10.5)
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
            run.font.bold = True
            run.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

        # 添加数据
        for i in range(t_rows) :
            t_row = t_table.add_row()
            # print("输出表格行: {}".format(in_table_data_list[i]))
            for j in range(t_cols) :
                # print("i=", i, "j=",j)
                # t_row.cells[j].text = in_table_data_list[i][j]
                t_row.cells[j].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # print(in_table_data_list[i][j])
                run = t_row.cells[j].paragraphs[0].add_run(in_table_data_list[i][j])
                run.font.name = '宋体'
                run.font.size = Pt(10.5)
                run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                run.font.bold = False
                run.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

                # t_row.cells[j].paragraphs[0].paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                # t_row.cells[j].vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    def pic(self, in_filename, in_title="", in_width=Cm(15), in_width_crop=1.0, in_height_crop=1.0):
        # 裁剪
        print("需要的pic文件为: {}".format(self.pic_filepath + in_filename + "_openid(" + self._id + ").png"))
        with Image.open(self.pic_filepath + in_filename + "_openid(" + self._id + ").png") as img:
            # 从图片(0,0)开始裁剪到(1/3横向长度，1/3纵向长度)
            img_new = img.crop((img.size[0] * (1-in_width_crop)/2.0, img.size[1] * (1-in_height_crop)/2.0, img.size[0] * (1-(1-in_width_crop)/2.0), img.size[1] * (1-(1-in_height_crop)/2.0)))

            t_temp_pic_filename = self.pic_filepath + "_temp_" + in_filename + "_openid(" + self._id + ").png"
            img_new.save(t_temp_pic_filename)
            self._temp_files.append(t_temp_pic_filename)
            print("临时的pic文件为: {}".format(t_temp_pic_filename))

        # 放入word
        para = self._doc.add_paragraph()
        para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        para.add_run().add_picture(t_temp_pic_filename, height=in_width * 1.4)

        # self._doc.add_picture(self.pic_path+"new_pic.png", width=in_width)

        if in_title!="" :
            self.para(in_title, in_style="Normal", in_center=True, in_indent=False, in_line_space=1.0)

    def save(self):
        self._doc.save(self.docx_filename)
        for i in self._temp_files :
            os.remove(i)
            print("删除临时文件: {}".format(i))

    # 分页符
    def page_break(self):
        self._doc.add_page_break()

    # 分节符
    def section_break(self, in_start_type=WD_SECTION_START.NEW_PAGE, in_rotate=False, in_is_linked_to_previous=True):
        sec = self._doc.add_section(start_type=in_start_type)
        # WD_SECTION_START.CONTINUOUS 1 连续分隔符
        # WD_SECTION_START.NEW_COLUMN 2 新列分隔符
        # WD_SECTION_START.NEW_PAGE   3 新页分隔符
        # WD_SECTION_START.EVEN_PAGE  4 偶页分隔符
        # WD_SECTION_START.ODD_PAGE   5 奇页分隔符
        if in_rotate :
            width = sec.page_width
            height = sec.page_height
            sec.page_width = height
            sec.page_height = width
            # sec.orientation = WD_ORIENTATION.LANDSCAPE  # 页面方向为横向
        # else:
        #     sec.orientation = WD_ORIENTATION.PORTRAIT   # 页面方向为纵向

        sec.footer.is_linked_to_previous = in_is_linked_to_previous
        return sec

    def get_docx_full_url(self):
        return self._docx_full_url


    # ==============添加页码的底层方式==============
    def _create_element(self, name):
        return OxmlElement(name)

    def _create_attribute(self, element, name, value):
        element.set(qn(name), value)

    def _run_add_str(self, in_run, in_str):
        t = self._create_element('w:t')
        self._create_attribute(t, 'xml:space', 'preserve')
        t.text = in_str
        in_run._r.append(t)

    def _run_add_page_index(self, in_run):
        fldChar1 = self._create_element('w:fldChar')
        self._create_attribute(fldChar1, 'w:fldCharType', 'begin')

        instrText = self._create_element('w:instrText')
        self._create_attribute(instrText, 'xml:space', 'preserve')
        instrText.text = "PAGE"

        fldChar2 = self._create_element('w:fldChar')
        self._create_attribute(fldChar2, 'w:fldCharType', 'end')

        in_run._r.append(fldChar1)
        in_run._r.append(instrText)
        in_run._r.append(fldChar2)

    def _run_add_total_page_num(self, in_run):
        fldChar3 = self._create_element('w:fldChar')
        self._create_attribute(fldChar3, 'w:fldCharType', 'begin')

        instrText2 = self._create_element('w:instrText')
        self._create_attribute(instrText2, 'xml:space', 'preserve')
        instrText2.text = "NUMPAGES"

        fldChar4 = self._create_element('w:fldChar')
        self._create_attribute(fldChar4, 'w:fldCharType', 'end')

        # in_run = paragraph.add_run()
        in_run._r.append(fldChar3)
        in_run._r.append(instrText2)
        in_run._r.append(fldChar4)

    def _add_page_number(self, paragraph, in_show_total=True):
        paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        page_run = paragraph.add_run()

        self._run_add_str(page_run, '第 ')
        self._run_add_page_index(page_run)          # 添加变量："PAGE"
        if in_show_total :
            self._run_add_str(page_run, ' 页/共 ')
            self._run_add_total_page_num(page_run)  # 添加变量："NUMPAGES"
            self._run_add_str(page_run, '页')
        else:
            self._run_add_str(page_run, ' 页')

    # 页脚添加页码
    def add_page_number(self, in_sec=0, in_show_total=True):
        if in_sec==0:
            in_sec = self._doc.sections[0]

        self._add_page_number(in_sec.footer.paragraphs[0], in_show_total=in_show_total)

    # 自定义样式
    def create_style(self, style_name, style_type, font_size=-1, font_bold=False, font_color=None, font_name=None, align=None):
        """
        创建一个样式
        :param align:
        :param _doc:
        :param style_name: 样式名称
        :param style_type: 样式类型，1：段落样式, 2：字符样式, 3：表格样式
        :param font_name:
        :param font_color:
        :param font_size:
        :return:
        """
        if font_color is None:
            font_color = []

        # 注意：必须要判断样式是否存在，否则重新添加会报错
        style_names = [style.name for style in self._doc.styles]
        if style_name in style_names:
            # print('样式已经存在，不需要重新添加！')
            return

        font_style = self._doc.styles.add_style(style_name, style_type)

        # 字体大小
        if font_size != -1:
            font_style.font.size = Pt(font_size)

        # 字体颜色
        # 比如：[0xff,0x00,0x00]
        if font_color and len(font_color) == 3:
            font_style.font.color.rgb = RGBColor(font_color[0], font_color[1], font_color[2])

        # 对齐方式
        # 注意：段落、表格才有对齐方式
        if style_type != 2 and align:
            font_style.paragraph_format.alignment = align
            # font_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            # font_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            # font_style.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT

        # 中文字体名称
        if font_name:
            font_style.font.name = font_name
            font_style._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

        font_style.font.bold = font_bold

        return font_style

    def get_current_section(self):
        t_len = len(self._doc.sections)
        return self._doc.sections[t_len-1]

def main():
    w=Word(in_filename="投资优化", in_id="o1Qen5J_nEK4hPTFQPq9Y6j_82hI")
    sec = w.section_break(in_start_type=WD_SECTION_START.NEW_PAGE, in_rotate=False, in_is_linked_to_previous=False)
    w.add_page_number(in_sec=sec, in_show_total=False)

    # for i in w._doc.styles :
    #     print(i)

    # 标题
    t_style = w.create_style(style_name="style2", style_type=1, font_size=30, font_color=[0xff, 0x00, 0x00])
    w.para("标题", in_style=t_style, in_center=True, in_indent=False)
    # w.para("标题", in_style="Subtitle", in_center=True, in_indent=False)

    w.heading("标题1", 1)
    w.heading("标题2", 2)
    w.heading("标题3", 3)
    w.heading("标题4", 4)
    w.heading("标题5", 5)
    w.para("正文", in_style="Normal")

    w.section_break(in_start_type=WD_SECTION_START.NEW_PAGE, in_rotate=False)
    w.table(
        in_table_head_list = ["序号", "项目", "规模", "投资", "备注"],
        in_table_data_list = [
            ["1", "投资优化", "3*240MVA", "1000万元", ""],
            ["2", "机组组合", "3*180MVA", "500万元", ""]
        ],
        in_title="计算结果汇总表",
    )
    # w.footer("落款", 1)


    w.pic(in_filename="NPS_IO_720h", in_title="NPS_IO_720h示意图", in_width_crop=0.9, in_height_crop=0.85)
    w.page_break()

    w.save()

if __name__ == "__main__" :
    main()
